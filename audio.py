import os
import streamlit as st
from dotenv import load_dotenv
import numpy as np
import torch
import pyaudio
import wave
from funasr import AutoModel
from sentence_transformers import SentenceTransformer
from openai import OpenAI
import json
import pandas as pd
import plotly.express as px
import re
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io

load_dotenv()
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
DEEPSEEK_API_BASE = "https://api.deepseek.com/v1"
client = OpenAI(api_key=DEEPSEEK_API_KEY , base_url="https://api.deepseek.com")

# 创建保存聊天记录的目录
CHAT_HISTORY_DIR = "chat_history"
CHART_IMAGES_DIR = "chart_images"  # 新增图表保存目录
if not os.path.exists(CHAT_HISTORY_DIR):
    os.makedirs(CHAT_HISTORY_DIR)
if not os.path.exists(CHART_IMAGES_DIR):  # 创建图表保存目录
    os.makedirs(CHART_IMAGES_DIR)

# 初始化session
if 'session_id' not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

def save_chat_history_to_word():
    """保存聊天记录到Word文档"""
    if st.session_state.chat_history:
        doc = Document()
        doc.add_heading('对话记录', 0)
        
        # 添加会话信息
        doc.add_paragraph(f'会话ID: {st.session_state.session_id}')
        doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('---' * 20)
        
        # 添加对话内容
        for i, record in enumerate(st.session_state.chat_history, 1):
            doc.add_heading(f'对话 {i}', level=1)
            doc.add_paragraph(f'时间: {record["timestamp"]}')
            doc.add_paragraph('问题: ' + record['query'])
            doc.add_paragraph('回答: ' + record['response']['answer'])
            
            # 如果有图表，添加图表说明和图片
            if record['response'].get('need_chart', False):
                chart_info = record['response']['chart']
                doc.add_paragraph('图表信息:')
                doc.add_paragraph(f'- 类型: {chart_info["type"]}')
                doc.add_paragraph(f'- 标题: {chart_info["title"]}')
                
                # 添加图片（如果存在）
                if "image_path" in record['response']:
                    image_path = record['response']["image_path"]
                    if os.path.exists(image_path):
                        doc.add_picture(image_path, width=Inches(6))
                    
                if "data" in chart_info:
                    doc.add_paragraph('- 数据:')
                    for item in chart_info["data"]:
                        doc.add_paragraph(f'  * {item["x"]}: {item["y"]}')
            
            doc.add_paragraph('---' * 20)
        
        # 保存到内存中
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
    return None

def save_chat_history():
    """保存聊天记录到JSON文件"""
    if st.session_state.chat_history:
        filename = f"{CHAT_HISTORY_DIR}/chat_{st.session_state.session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump({
                'session_id': st.session_state.session_id,
                'timestamp': datetime.now().isoformat(),
                'history': st.session_state.chat_history
            }, f, ensure_ascii=False, indent=2)
        return filename
    return None

device = "cpu"

# 初始化 FunASR 模型
model = AutoModel(model="paraformer-zh")

# 录音参数
CHUNK = 1024
FORMAT = pyaudio.paInt16
CHANNELS = 1
RATE = 16000  # FunASR 默认 16kHz
RECORD_SECONDS = 5  # 录音时长（可调整）
WAVE_OUTPUT_FILENAME = "temp_audio.wav"

# 初始化 PyAudio
audio = pyaudio.PyAudio()

# 开始录音
def record_and_transcribe(duration=5, fs=16000):
    st.info("开始录音，请讲话...")
    
    stream = audio.open(format=FORMAT,
                        channels=CHANNELS,
                        rate=RATE,
                        input=True,
                        frames_per_buffer=CHUNK)

    frames = []
    for i in range(0, int(RATE / CHUNK * duration)):
        data = stream.read(CHUNK)
        frames.append(data)

    stream.stop_stream()
    stream.close()

    # 保存录音为 WAV 文件
    wf = wave.open(WAVE_OUTPUT_FILENAME, 'wb')
    wf.setnchannels(CHANNELS)
    wf.setsampwidth(audio.get_sample_size(FORMAT))
    wf.setframerate(RATE)
    wf.writeframes(b''.join(frames))
    wf.close()

    st.success("录音完成，正在识别...")
    # 使用 FunASR 识别
    result = model.generate(input=WAVE_OUTPUT_FILENAME)
    return result[0]["text"]

# 本地Embedding
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer('BAAI/bge-large-zh-v1.5', device=device)

embedding_model = load_embedding_model()
print("embedding_model loaded")
def get_embedding(text):
    embedding = embedding_model.encode(text, normalize_embeddings=True)
    return np.array(embedding, dtype='float32')


# 读取完整文件
def read_full_file(filename="demo.txt"):
    """读取documents目录下的完整文件内容"""
    file_path = os.path.join("documents", filename)
    if os.path.exists(file_path):
        with open(file_path, "r") as f:
            return f.read()
    else:
        print(f"文件 {filename} 不存在")
        return f"文件 {filename} 不存在"

# DeepSeek Chat推理
# DeepSeek Chat API调用
def deepseek_chat_json(prompt):
    # 创建一个空的占位符用于显示生成过程
    response_placeholder = st.empty()
    accumulated_content = ""
    
    # 使用stream=True来启用流式输出
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
        stream=True
    )
    
    # 逐步接收并显示内容
    for chunk in response:
        if chunk.choices[0].delta.content is not None:
            content_piece = chunk.choices[0].delta.content
            accumulated_content += content_piece
            
            response_placeholder.markdown("**正在生成回答：**\n" + accumulated_content)
    
    content = accumulated_content
    print("原始响应内容:", content)
    
    # 尝试直接解析完整的JSON
    try:
        result = json.loads(content)
        print("成功解析JSON:", result)
        return result
    except json.JSONDecodeError as e:
        print(f"JSON解析失败: {e}")
        
        # 尝试查找JSON内容
        json_pattern = r'```json\s*([\s\S]*?)\s*```'
        json_match = re.search(json_pattern, content)
        if json_match:
            try:
                json_content = json_match.group(1)
                print("找到JSON块:", json_content)
                result = json.loads(json_content)
                print("成功解析JSON块:", result)
                return result
            except json.JSONDecodeError as e2:
                print(f"JSON块解析失败: {e2}")
        
        # 查找不带代码块标记的JSON
        json_pattern2 = r'(\{[\s\S]*\})'
        json_match2 = re.search(json_pattern2, content)
        if json_match2:
            try:
                json_content2 = json_match2.group(1)
                print("找到JSON对象:", json_content2)
                result = json.loads(json_content2)
                print("成功解析JSON对象:", result)
                return result
            except json.JSONDecodeError as e3:
                print(f"JSON对象解析失败: {e3}")
        
        # 所有尝试都失败，返回默认对象
        print("所有JSON解析尝试失败，返回默认对象")
        return {"answer": content, "need_chart": False, "chart": {}}

# Streamlit UI
st.title("🎙️语音识别+DeepSeek图表生成Demo")

# 显示当前会话ID和导出按钮
col1, col2 = st.columns([3, 1])
with col1:
    st.text(f"会话ID: {st.session_state.session_id}")
with col2:
        # 导出Word按钮
    if st.button("📄 导出Word"):
        doc_io = save_chat_history_to_word()
        if doc_io:
            st.session_state["doc_io"] = doc_io  # 存进session_state
            st.success("Word文档已生成，请点击下方的下载按钮")

# 如果session_state中有doc_io，就显示下载按钮
if "doc_io" in st.session_state and st.session_state["doc_io"]:
    st.download_button(
        label="⬇️ 下载Word文档",
        data=st.session_state["doc_io"].getvalue(),
        file_name=f"对话记录_{st.session_state.session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# 显示当前对话历史记录
if st.session_state.chat_history:
    st.write("---")
    st.write("### 当前对话历史记录")
    for i, chat in enumerate(st.session_state.chat_history, 1):
        st.write(f"**对话 {i}:**")
        st.write(f"时间: {chat['timestamp']}")
        st.write(f"问题: {chat['query']}")
        st.write(f"回答: {chat['response']['answer']}")
        st.write("---")

st.write("### 开始新的对话")
if st.button("🎤 点击开始录音识别"):
    query_text = record_and_transcribe(duration=10)
    st.write(f"**识别结果：** {query_text}")

    with st.spinner('本地知识库召回中...'):
        context = read_full_file("demo1.txt")

    prompt = f"""
    根据给定的财务数据回答用户问题，并判断回答是否适合图表展示。

    财务数据：
    {context}

    用户的问题：{query_text} 

    请严格使用以下JSON格式回复：
    
```json
    {{
      "answer": "完整的文字回答",
      "need_chart": true 或 false,
      "chart": {{
        "type": "bar,line,pie,stacked_bar,horizontal_bar",
        "title": "图表标题",
        "x_label": "X轴名称",
        "y_label": "Y轴名称",
        "data": [
          {{"x": "X轴项", "y": 数值}}
        ]
      }}
    }}
```

    说明1：如果问题不适合图表展示，请设need_chart为false并留空chart。
    说明2: bar,line,pie,stacked_bar,horizontal_bar分别为柱状图、折线图、饼图、堆叠图、水平柱状图，如果用户没有要求，你可以自由选择合适的类型。
    说明3：由于用户问题是通过语音识别的，可能存在一些错误，请根据实际情况进行调整.
    """
    
    with st.spinner('DeepSeek 生成答案中 ...'):
        response = deepseek_chat_json(prompt)
        
        # 保存对话记录
        chat_record = {
            'timestamp': datetime.now().isoformat(),
            'query': query_text,
            'response': response
        }
        st.session_state.chat_history.append(chat_record)
        
        # 显示回答
        if response.get('need_chart', False):
            chart = response["chart"]
            chart_df = pd.DataFrame(chart["data"])

            st.subheader(chart.get("title", "数据图表"))
            chart_type = chart.get("type", "bar")

            # 生成图表并保存
            fig = None
            if chart_type == "bar":
                fig = px.bar(chart_df, x='x', y='y', title=chart.get("title"))
            elif chart_type == "line":
                fig = px.line(chart_df, x='x', y='y', title=chart.get("title"))
            elif chart_type == "pie":
                fig = px.pie(chart_df, names="x", values="y", title=chart.get("title"))
            elif chart_type == "stacked_bar":
                fig = px.bar(chart_df, x='x', y='y', title=chart.get("title"))
            elif chart_type == "horizontal_bar":
                fig = px.bar(chart_df, x='y', y='x', orientation='h', title=chart.get("title"))
            
            if fig:
                # 保存图片
                image_filename = f"chart_{st.session_state.session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                image_path = os.path.join(CHART_IMAGES_DIR, image_filename)
                fig.write_image(image_path)
                
                # 在response中添加图片路径
                response["image_path"] = image_path
                
                # 显示图表
                st.plotly_chart(fig)
            else:
                st.warning("暂不支持此图表类型")
        else:
            st.info("该问题无需图表展示。")

        st.success("对话已完成")
